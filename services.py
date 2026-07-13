"""
app/services.py

All business logic — async. FAISS completely replaced with Pinecone.

Flow:
  Upload  → extract text → chunk → embed → upsert to Pinecone cloud
  Ask     → embed question → Pinecone similarity search → Gemini answer

Where FAISS was used (OLD):              Where Pinecone is used (NEW):
  _build_faiss_index()              →      _upsert_to_pinecone()
  app_state.faiss_index             →      app_state.pinecone_index
  faiss_index.similarity_search()   →      pinecone_index.query()
  saved to disk (faiss_index/)      →      stored in Pinecone cloud
"""

import asyncio
import uuid
from io import BytesIO
from typing import List

from fastapi import UploadFile
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain.chains.question_answering import load_qa_chain
from langchain.chains import LLMChain
from langchain.prompts import PromptTemplate
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.docstore.document import Document as LCDocument
from PyPDF2 import PdfReader
from docx import Document

from app.state import app_state


# ─────────────────────────────────────────
# STEP 1 — File Text Extraction
# ─────────────────────────────────────────
def _extract_text_from_file(file_bytes: bytes, filename: str) -> str:
    """
    Extract raw text from PDF, TXT, or DOCX.
    Same logic as your original Streamlit version.
    """
    ext = filename.split('.')[-1].lower()
    text = ""

    if ext == 'pdf':
        reader = PdfReader(BytesIO(file_bytes))
        for page in reader.pages:
            text += page.extract_text() or ""

    elif ext == 'txt':
        text = file_bytes.decode("utf-8")

    elif ext == 'docx':
        doc = Document(BytesIO(file_bytes))
        for para in doc.paragraphs:
            text += para.text + "\n"

    return text


# ─────────────────────────────────────────
# STEP 2 — Chunking
# ─────────────────────────────────────────
def _chunk_text(text: str) -> list:
    """
    Split large text into overlapping chunks.
    chunk_size=10000, chunk_overlap=1000 — same as your Streamlit version.
    Each chunk becomes one vector in Pinecone.
    """
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=10000,
        chunk_overlap=1000
    )
    return splitter.split_text(text)


# ─────────────────────────────────────────
# STEP 3 — Embed + Upsert to Pinecone
# (replaces _build_faiss_index)
# ─────────────────────────────────────────
def _upsert_to_pinecone(chunks: list):
    """
    Embed each chunk → upsert vector to Pinecone cloud.

    OLD (FAISS):
        index = FAISS.from_texts(chunks, embedding=app_state.embeddings)
        index.save_local("faiss_index")          ← saved to disk, lost on restart
        app_state.faiss_index = index            ← only in local memory

    NEW (Pinecone):
        embed each chunk → upsert to Pinecone cloud
        ← persistent across restarts
        ← accessible from any server instance
        ← supports concurrent reads natively
    """
    vectors = []
    for chunk in chunks:
        vector = app_state.embeddings.embed_query(chunk)  # 384-dim float list
        vectors.append({
            "id": str(uuid.uuid4()),          # unique ID per chunk
            "values": vector,                  # the actual embedding
            "metadata": {"text": chunk}        # store original text for retrieval
        })

    # Upsert in batches of 100 (Pinecone recommended)
    batch_size = 100
    for i in range(0, len(vectors), batch_size):
        app_state.pinecone_index.upsert(vectors=vectors[i:i + batch_size])

    print(f"  ✅ Upserted {len(vectors)} vectors to Pinecone.")


# ─────────────────────────────────────────
# STEP 4 — Similarity Search from Pinecone
# (replaces faiss_index.similarity_search)
# ─────────────────────────────────────────
def _search_pinecone(question: str, top_k: int = 4) -> list:
    """
    Embed the question → query Pinecone → return top-k matching chunks.

    OLD (FAISS):
        docs = app_state.faiss_index.similarity_search(question)
        ← reads from local memory / disk

    NEW (Pinecone):
        query_vector = embed(question)
        results = pinecone_index.query(vector=query_vector, top_k=4)
        ← reads from Pinecone cloud
        ← stateless, concurrent-safe, no disk I/O
    """
    # Embed the user's question
    query_vector = app_state.embeddings.embed_query(question)

    # Query Pinecone for top-k similar vectors
    results = app_state.pinecone_index.query(
        vector=query_vector,
        top_k=top_k,
        include_metadata=True    # need metadata to get original text back
    )

    # Convert Pinecone results → LangChain Document objects
    # (QA chain expects LangChain Documents)
    docs = [
        LCDocument(page_content=match["metadata"]["text"])
        for match in results["matches"]
        if "text" in match.get("metadata", {})
    ]
    return docs


# ─────────────────────────────────────────
# Document Processing (Upload Handler)
# ─────────────────────────────────────────
async def process_uploaded_files(files: List[UploadFile]):
    """
    Full pipeline:
    Read files → extract text → chunk → embed → upsert to Pinecone

    run_in_executor: offloads CPU-bound embedding work to thread pool
    so it doesn't block the async event loop for other users.
    """
    raw_text = ""
    for file in files:
        file_bytes = await file.read()
        raw_text += _extract_text_from_file(file_bytes, file.filename) + "\n"

    chunks = _chunk_text(raw_text)

    loop = asyncio.get_event_loop()
    await loop.run_in_executor(None, _upsert_to_pinecone, chunks)


# ─────────────────────────────────────────
# Document QA
# ─────────────────────────────────────────
def _get_qa_chain():
    prompt_template = """
    Answer the question as detailed as possible from the provided context.
    Make sure to provide all the details.
    If the answer is not in the context, say: "Answer is not available in the context."
    Do not make up answers.

    Context:\n {context}?\n
    Question: \n{question}\n

    Answer:
    """
    model = ChatGoogleGenerativeAI(model="gemini-2.5-flash", temperature=0.3)
    prompt = PromptTemplate(template=prompt_template, input_variables=["context", "question"])
    return load_qa_chain(model, chain_type="stuff", prompt=prompt)


async def answer_from_documents(question: str) -> str:
    """
    1. Search Pinecone for relevant chunks
    2. Pass to Gemini QA chain
    Both offloaded to thread pool — non-blocking.
    """
    loop = asyncio.get_event_loop()

    # Step 1: Pinecone search (replaces faiss_index.similarity_search)
    docs = await loop.run_in_executor(
        None, lambda: _search_pinecone(question)
    )

    if not docs:
        return "No relevant content found. Please upload documents first."

    # Step 2: Gemini answer generation
    chain = _get_qa_chain()
    response = await loop.run_in_executor(
        None,
        lambda: chain(
            {"input_documents": docs, "question": question},
            return_only_outputs=True
        )
    )
    return response["output_text"]


# ─────────────────────────────────────────
# Summarizer
# ─────────────────────────────────────────
def _get_summarize_chain():
    prompt_template = """
    Summarize the following text in a clear and concise way:

    Text:
    {text}

    Summary:
    """
    model = ChatGoogleGenerativeAI(model="gemini-2.5-flash", temperature=0.3)
    prompt = PromptTemplate(template=prompt_template, input_variables=["text"])
    return LLMChain(llm=model, prompt=prompt)


async def summarize_text(text: str) -> str:
    chain = _get_summarize_chain()
    loop = asyncio.get_event_loop()
    return await loop.run_in_executor(None, lambda: chain.run(text=text))


# ─────────────────────────────────────────
# Career Counseling
# ─────────────────────────────────────────
def _get_career_chain():
    prompt_template = """
    You are a helpful and professional career counselor.
    Answer the following query with actionable advice and clarity.

    Question:
    {question}

    Answer:
    """
    model = ChatGoogleGenerativeAI(model="gemini-2.5-flash", temperature=0.5)
    prompt = PromptTemplate(template=prompt_template, input_variables=["question"])
    return LLMChain(llm=model, prompt=prompt)


async def career_counseling_response(question: str) -> str:
    chain = _get_career_chain()
    loop = asyncio.get_event_loop()
    return await loop.run_in_executor(None, lambda: chain.run(question=question))


# ─────────────────────────────────────────
# Personal & Emotional Support
# ─────────────────────────────────────────
def _get_personal_chain():
    prompt_template = """
    You are a friendly, understanding, and supportive AI assistant.
    Respond with empathy, encouragement, and practical advice.

    Question:
    {question}

    Answer:
    """
    model = ChatGoogleGenerativeAI(model="gemini-2.5-flash", temperature=0.7)
    prompt = PromptTemplate(template=prompt_template, input_variables=["question"])
    return LLMChain(llm=model, prompt=prompt)


async def personal_support_response(question: str) -> str:
    chain = _get_personal_chain()
    loop = asyncio.get_event_loop()
    return await loop.run_in_executor(None, lambda: chain.run(question=question))


# ─────────────────────────────────────────
# Entertainment
# ─────────────────────────────────────────
def _get_entertainment_chain():
    prompt_template = """
    You are a friendly AI assistant for entertainment recommendations.
    Help with movies, shows, songs, web series, and genre-based suggestions.

    Question:
    {question}

    Answer:
    """
    model = ChatGoogleGenerativeAI(model="gemini-2.5-flash", temperature=0.7)
    prompt = PromptTemplate(template=prompt_template, input_variables=["question"])
    return LLMChain(llm=model, prompt=prompt)


async def entertainment_response(question: str) -> str:
    chain = _get_entertainment_chain()
    loop = asyncio.get_event_loop()
    return await loop.run_in_executor(None, lambda: chain.run(question=question))


# ─────────────────────────────────────────
# MCQ & Notes Generator
# ─────────────────────────────────────────
def _get_mcq_chain():
    prompt_template = """
    You are a helpful educational assistant. From the text below:

    1. Generate 3 MCQs with 4 options each and indicate the correct answer.
    2. Provide short notes summarizing the key points.

    Text:
    {text}

    Output:
    """
    model = ChatGoogleGenerativeAI(model="gemini-2.5-flash", temperature=0.4)
    prompt = PromptTemplate(template=prompt_template, input_variables=["text"])
    return LLMChain(llm=model, prompt=prompt)


async def generate_mcq_and_notes(text: str) -> str:
    chain = _get_mcq_chain()
    loop = asyncio.get_event_loop()
    return await loop.run_in_executor(None, lambda: chain.run(text=text))
